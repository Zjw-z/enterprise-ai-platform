import asyncio
import io
import json
import zipfile

import httpx

from app.knowledge import (
    DocumentParseError,
    DocumentQualityGate,
    FallbackDocumentParser,
    KnowledgeIngestionService,
    KnowledgeService,
    MinerUPrecisionParser,
    NativeDocumentParser,
    ParsedBlock,
    ParsedDocument,
)
from app.system.database import SystemDatabase
from app.vector import VectorOutboxService


def test_quality_gate_rejects_empty_and_corrupted_documents() -> None:
    gate = DocumentQualityGate(minimum_score=60)
    empty = ParsedDocument(
        filename="scan.pdf",
        parser="native",
        blocks=(ParsedBlock(""),),
    )
    report = gate.inspect(empty)
    assert report.passed is False
    assert report.score < 60
    assert report.issues[0].code == "insufficient_text"

    corrupted = ParsedDocument(
        filename="broken.txt",
        parser="native",
        blocks=(ParsedBlock("有效文本" * 10 + "\ufffd" * 20),),
    )
    report = gate.inspect(corrupted)
    assert report.passed is False
    assert any(
        issue.code == "encoding_corruption"
        for issue in report.issues
    )


def test_native_docx_preserves_headings_tables_and_paragraphs() -> None:
    from docx import Document

    stream = io.BytesIO()
    source = Document()
    source.add_heading("差旅制度", level=1)
    source.add_paragraph("员工出差前必须审批。")
    table = source.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "城市"
    table.cell(0, 1).text = "限额"
    source.save(stream)

    result = asyncio.run(
        NativeDocumentParser().parse(
            filename="policy.docx",
            content=stream.getvalue(),
            content_type=(
                "application/vnd.openxmlformats-officedocument."
                "wordprocessingml.document"
            ),
        )
    )
    assert [block.kind for block in result.blocks] == [
        "heading",
        "paragraph",
        "table",
    ]
    assert result.blocks[0].heading == "差旅制度"
    assert "城市 | 限额" in result.blocks[2].text


def test_mineru_precision_adapter_uploads_polls_and_reads_structure() -> None:
    archive_stream = io.BytesIO()
    with zipfile.ZipFile(archive_stream, "w") as archive:
        archive.writestr(
            "demo_content_list.json",
            json.dumps(
                [
                    {
                        "type": "text",
                        "text": "第一段制度正文",
                        "page_idx": 0,
                        "bbox": [1, 2, 3, 4],
                    },
                    {
                        "type": "table",
                        "table_body": "|城市|限额|",
                        "page_idx": 1,
                    },
                ],
                ensure_ascii=False,
            ),
        )
    polls = 0

    def handler(request: httpx.Request) -> httpx.Response:
        nonlocal polls
        if request.url.path == "/api/v4/file-urls/batch":
            return httpx.Response(
                200,
                json={
                    "code": 0,
                    "data": {
                        "batch_id": "batch-1",
                        "file_urls": ["https://upload.example/file"],
                    },
                },
            )
        if request.url.host == "upload.example":
            assert request.method == "PUT"
            return httpx.Response(200)
        if request.url.path.endswith(
            "/api/v4/extract-results/batch/batch-1"
        ):
            polls += 1
            return httpx.Response(
                200,
                json={
                    "code": 0,
                    "data": {
                        "extract_result": [
                            (
                                {"state": "running"}
                                if polls == 1
                                else {
                                    "state": "done",
                                    "full_zip_url": (
                                        "https://cdn.example/result.zip"
                                    ),
                                }
                            )
                        ]
                    },
                },
            )
        if request.url.host == "cdn.example":
            return httpx.Response(200, content=archive_stream.getvalue())
        return httpx.Response(404)

    async def scenario() -> ParsedDocument:
        async with httpx.AsyncClient(
            transport=httpx.MockTransport(handler)
        ) as client:
            parser = MinerUPrecisionParser(
                base_url="https://mineru.net",
                api_token="test-token",
                poll_interval_seconds=0.001,
                timeout_seconds=10,
                client=client,
            )
            return await parser.parse(
                filename="demo.pdf",
                content=b"pdf",
                content_type="application/pdf",
            )

    result = asyncio.run(scenario())
    assert result.parser == "mineru_api"
    assert result.page_count == 2
    assert result.blocks[0].page == 1
    assert result.blocks[1].kind == "table"


def test_parser_fallback_records_primary_failure() -> None:
    class BrokenParser:
        name = "broken"

        async def parse(self, **_: object) -> ParsedDocument:
            raise DocumentParseError("remote unavailable")

    result = asyncio.run(
        FallbackDocumentParser(
            BrokenParser(), NativeDocumentParser()
        ).parse(
            filename="notes.txt",
            content="可用于知识库的有效文本。".encode(),
            content_type="text/plain",
        )
    )
    assert result.parser == "native"
    assert result.metadata["fallback_from"] == "broken"
    assert "remote unavailable" in result.metadata["fallback_reason"]


def test_streaming_batch_persists_each_independent_outcome() -> None:
    class FakeObjectStore:
        def __init__(self) -> None:
            self.objects: dict[str, bytes] = {}

        async def put(
            self,
            *,
            object_key: str,
            content: bytes,
            content_type: str,
        ) -> None:
            self.objects[object_key] = content

        async def delete(self, object_key: str) -> None:
            self.objects.pop(object_key, None)

        async def get(self, object_key: str) -> bytes:
            return self.objects[object_key]

        async def presigned_put(
            self, object_key: str, *, expires_seconds: int
        ) -> str:
            return f"http://minio.test/{object_key}?expires={expires_seconds}"

        async def stat(self, object_key: str) -> dict[str, object]:
            content = self.objects[object_key]
            return {
                "size": len(content),
                "content_type": "text/plain",
                "etag": "fake-etag",
            }

    class MixedParser:
        async def parse(
            self,
            *,
            filename: str,
            content: bytes,
            content_type: str,
        ) -> ParsedDocument:
            if filename == "broken.txt":
                raise DocumentParseError("模拟解析器异常")
            text = (
                "有效的企业制度正文，内容足够进入向量索引。"
                if filename == "good.txt"
                else ""
            )
            return ParsedDocument(
                filename=filename,
                parser="mixed",
                blocks=(ParsedBlock(text),),
            )

    async def scenario() -> None:
        database = SystemDatabase("sqlite+aiosqlite:///:memory:")
        await database.initialize()
        knowledge = KnowledgeService(
            database,
            VectorOutboxService(database),
            collection_name="knowledge",
            embedding_model="test",
            embedding_dimensions=3,
        )
        base = await knowledge.create_base(
            tenant_id="tenant-a",
            name="batch-test",
            description="",
            visibility="tenant",
            allowed_roles=[],
            actor_id="admin",
        )
        object_store = FakeObjectStore()
        ingestion = KnowledgeIngestionService(
            knowledge,
            object_store,  # type: ignore[arg-type]
            parser=MixedParser(),
            quality_gate=DocumentQualityGate(
                minimum_score=60,
                minimum_characters=20,
            ),
        )
        async def stream_files():
            for item in (
                {
                    "filename": "good.txt",
                    "content_type": "text/plain",
                    "content": b"good",
                },
                {
                    "filename": "broken.txt",
                    "content_type": "text/plain",
                    "content": b"broken",
                },
                {
                    "filename": "empty.txt",
                    "content_type": "text/plain",
                    "content": b"empty",
                },
                {
                    "filename": "oversized.pdf",
                    "upload_error": "文件超过单文件大小限制。",
                },
            ):
                yield item

        batch = await ingestion.ingest_batch(
            tenant_id="tenant-a",
            knowledge_base_id=base["id"],
            actor_id="admin",
            files=stream_files(),
            total_count=4,
        )
        assert batch["status"] == "partial_failed"
        assert batch["success_count"] == 1
        assert batch["failed_count"] == 2
        assert batch["quality_failed_count"] == 1
        assert next(
            item
            for item in batch["items"]
            if item["title"] == "oversized.pdf"
        )["parsing_error"] == "文件超过单文件大小限制。"
        documents = await knowledge.list_documents(
            tenant_id="tenant-a",
            knowledge_base_id=base["id"],
        )
        assert {item["parsing_status"] for item in documents} == {
            "completed",
            "failed",
            "quality_failed",
        }
        assert next(
            item
            for item in documents
            if item["parsing_status"] == "failed"
        )["parsing_error"] == "模拟解析器异常"
        batches = await knowledge.list_ingestion_batches(
            tenant_id="tenant-a",
            knowledge_base_id=base["id"],
        )
        assert batches[0]["total_count"] == 4

        queued = await ingestion.submit(
            tenant_id="tenant-a",
            knowledge_base_id=base["id"],
            filename="good.txt",
            content_type="text/plain",
            content=b"queued-good",
            actor_id="admin",
        )
        assert queued["parsing_status"] == "pending"
        assert await ingestion.process_once() == 1
        processed = await knowledge.get_document(
            tenant_id="tenant-a",
            document_id=queued["id"],
        )
        assert processed["parsing_status"] == "completed"
        assert processed["parsing_attempts"] == 1

        intent = await ingestion.create_upload_intent(
            tenant_id="tenant-a",
            knowledge_base_id=base["id"],
            filename="good.txt",
            content_type="text/plain",
            actor_id="admin",
            expires_seconds=900,
        )
        direct_document = intent["document"]
        assert direct_document["parsing_status"] == "uploading"
        object_store.objects[direct_document["object_key"]] = (
            b"direct-upload"
        )
        committed = await ingestion.commit_upload(
            tenant_id="tenant-a",
            knowledge_base_id=base["id"],
            document_id=direct_document["id"],
            maximum_bytes=1024,
        )
        assert committed["parsing_status"] == "pending"
        assert await ingestion.process_once() == 1
        direct_processed = await knowledge.get_document(
            tenant_id="tenant-a",
            document_id=direct_document["id"],
        )
        assert direct_processed["parsing_status"] == "completed"
        assert direct_processed["content_hash"] != "pending"
        await database.close()

    asyncio.run(scenario())
