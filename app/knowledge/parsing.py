"""Enterprise document parsing adapters and ingestion quality gate."""

from __future__ import annotations

import asyncio
import io
import json
import re
import time
import zipfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Protocol

import httpx


class DocumentParseError(RuntimeError):
    """A stable, user-safe parsing failure."""


@dataclass(frozen=True)
class ParsedBlock:
    """A source-addressable unit produced by a parser."""

    text: str
    kind: str = "text"
    page: int | None = None
    heading: str | None = None
    metadata: dict[str, Any] = field(default_factory=dict)


@dataclass(frozen=True)
class ParsedDocument:
    """Canonical parser output consumed by chunking and quality checks."""

    filename: str
    parser: str
    blocks: tuple[ParsedBlock, ...]
    page_count: int | None = None
    metadata: dict[str, Any] = field(default_factory=dict)

    @property
    def text(self) -> str:
        return "\n\n".join(block.text for block in self.blocks if block.text)


class DocumentParser(Protocol):
    """The parsing seam: adapters hide transport and format complexity."""

    async def parse(
        self,
        *,
        filename: str,
        content: bytes,
        content_type: str,
    ) -> ParsedDocument: ...


class NativeDocumentParser:
    """Offline fallback for text, HTML, PDF and DOCX documents."""

    name = "native"

    async def parse(
        self,
        *,
        filename: str,
        content: bytes,
        content_type: str,
    ) -> ParsedDocument:
        return await asyncio.to_thread(
            self._parse_sync, filename, content
        )

    def _parse_sync(
        self, filename: str, content: bytes
    ) -> ParsedDocument:
        suffix = Path(filename).suffix.lower()
        blocks: list[ParsedBlock] = []
        page_count: int | None = None
        if suffix in {".txt", ".md", ".markdown", ".csv"}:
            blocks.append(
                ParsedBlock(content.decode("utf-8-sig"), kind="text")
            )
        elif suffix in {".html", ".htm"}:
            raw = content.decode("utf-8-sig")
            text = re.sub(r"\s+", " ", re.sub(r"<[^>]+>", " ", raw))
            blocks.append(ParsedBlock(text, kind="html"))
        elif suffix == ".pdf":
            from pypdf import PdfReader

            reader = PdfReader(io.BytesIO(content))
            page_count = len(reader.pages)
            blocks.extend(
                ParsedBlock(
                    page.extract_text() or "",
                    kind="page",
                    page=index,
                )
                for index, page in enumerate(reader.pages, start=1)
            )
        elif suffix == ".docx":
            from docx import Document

            document = Document(io.BytesIO(content))
            for paragraph in document.paragraphs:
                text = paragraph.text.strip()
                if text:
                    style = paragraph.style.name.lower()
                    blocks.append(
                        ParsedBlock(
                            text,
                            kind=(
                                "heading"
                                if style.startswith("heading")
                                else "paragraph"
                            ),
                            heading=text if style.startswith("heading") else None,
                        )
                    )
            for table_index, table in enumerate(document.tables, start=1):
                rows = [
                    " | ".join(cell.text.strip() for cell in row.cells)
                    for row in table.rows
                ]
                if rows:
                    blocks.append(
                        ParsedBlock(
                            "\n".join(rows),
                            kind="table",
                            metadata={"table_index": table_index},
                        )
                    )
        else:
            raise DocumentParseError(
                f"不支持的文档格式：{suffix or '无扩展名'}"
            )
        return ParsedDocument(
            filename=filename,
            parser=self.name,
            blocks=tuple(blocks),
            page_count=page_count,
        )


class MinerUPrecisionParser:
    """MinerU precision API adapter for local-file upload and polling."""

    name = "mineru_api"

    def __init__(
        self,
        *,
        base_url: str,
        api_token: str,
        model_version: str = "vlm",
        language: str = "ch",
        enable_table: bool = True,
        enable_formula: bool = True,
        poll_interval_seconds: float = 2.0,
        timeout_seconds: float = 300.0,
        client: httpx.AsyncClient | None = None,
    ) -> None:
        if not api_token:
            raise ValueError("MinerU API Token 未配置。")
        self.base_url = base_url.rstrip("/")
        self.api_token = api_token
        self.model_version = model_version
        self.language = language
        self.enable_table = enable_table
        self.enable_formula = enable_formula
        self.poll_interval_seconds = poll_interval_seconds
        self.timeout_seconds = timeout_seconds
        self._client = client

    async def parse(
        self,
        *,
        filename: str,
        content: bytes,
        content_type: str,
    ) -> ParsedDocument:
        owns_client = self._client is None
        client = self._client or httpx.AsyncClient(
            timeout=httpx.Timeout(60.0)
        )
        try:
            headers = {"Authorization": f"Bearer {self.api_token}"}
            request = await client.post(
                f"{self.base_url}/api/v4/file-urls/batch",
                headers=headers,
                json={
                    "files": [{"name": Path(filename).name}],
                    "model_version": self.model_version,
                    "language": self.language,
                    "enable_table": self.enable_table,
                    "enable_formula": self.enable_formula,
                },
            )
            payload = self._response_payload(request, "申请上传地址")
            data = payload.get("data") or {}
            batch_id = str(data.get("batch_id") or "")
            upload_urls = data.get("file_urls") or []
            if not batch_id or not upload_urls:
                raise DocumentParseError("MinerU 未返回任务或上传地址。")
            upload = await client.put(upload_urls[0], content=content)
            upload.raise_for_status()
            result_url = (
                f"{self.base_url}/api/v4/extract-results/batch/{batch_id}"
            )
            deadline = time.monotonic() + self.timeout_seconds
            while time.monotonic() < deadline:
                response = await client.get(result_url, headers=headers)
                status_payload = self._response_payload(
                    response, "查询解析结果"
                )
                results = (
                    (status_payload.get("data") or {}).get("extract_result")
                    or []
                )
                if results:
                    result = results[0]
                    state = result.get("state")
                    if state == "done":
                        zip_url = result.get("full_zip_url")
                        if not zip_url:
                            raise DocumentParseError(
                                "MinerU 解析成功但未返回结果文件。"
                            )
                        archive = await client.get(zip_url)
                        archive.raise_for_status()
                        return self._read_archive(filename, archive.content)
                    if state == "failed":
                        raise DocumentParseError(
                            "MinerU 解析失败："
                            + str(result.get("err_msg") or "未知原因")
                        )
                await asyncio.sleep(self.poll_interval_seconds)
            raise DocumentParseError(
                f"MinerU 解析超过 {self.timeout_seconds:g} 秒。"
            )
        except httpx.HTTPError as exc:
            raise DocumentParseError(
                f"MinerU 网络请求失败：{exc}"
            ) from exc
        finally:
            if owns_client:
                await client.aclose()

    @staticmethod
    def _response_payload(
        response: httpx.Response, operation: str
    ) -> dict[str, Any]:
        response.raise_for_status()
        payload = response.json()
        if payload.get("code") != 0:
            raise DocumentParseError(
                f"MinerU {operation}失败："
                f"{payload.get('msg') or '未知原因'}"
            )
        return payload

    def _read_archive(
        self, filename: str, archive_content: bytes
    ) -> ParsedDocument:
        try:
            with zipfile.ZipFile(io.BytesIO(archive_content)) as archive:
                names = archive.namelist()
                content_name = next(
                    (
                        name
                        for name in names
                        if name.endswith("_content_list.json")
                    ),
                    None,
                )
                if content_name:
                    items = json.loads(
                        archive.read(content_name).decode("utf-8")
                    )
                    blocks = tuple(
                        block
                        for item in items
                        if (block := self._content_block(item)) is not None
                    )
                    if blocks:
                        return ParsedDocument(
                            filename=filename,
                            parser=self.name,
                            blocks=blocks,
                            page_count=self._page_count(items),
                            metadata={"model_version": self.model_version},
                        )
                markdown_name = next(
                    (name for name in names if name.endswith(".md")),
                    None,
                )
                if not markdown_name:
                    raise DocumentParseError(
                        "MinerU 结果中没有 Markdown 或结构化内容。"
                    )
                markdown = archive.read(markdown_name).decode("utf-8")
                return ParsedDocument(
                    filename=filename,
                    parser=self.name,
                    blocks=(ParsedBlock(markdown, kind="markdown"),),
                    metadata={"model_version": self.model_version},
                )
        except (zipfile.BadZipFile, UnicodeDecodeError, json.JSONDecodeError) as exc:
            raise DocumentParseError(
                "MinerU 返回的结果包无法读取。"
            ) from exc

    @staticmethod
    def _content_block(item: dict[str, Any]) -> ParsedBlock | None:
        kind = str(item.get("type") or "text")
        text = str(
            item.get("text")
            or item.get("content")
            or item.get("img_caption")
            or ""
        ).strip()
        if not text and kind == "table":
            text = str(item.get("table_body") or "").strip()
        if not text:
            return None
        page = item.get("page_idx")
        return ParsedBlock(
            text=text,
            kind=kind,
            page=int(page) + 1 if isinstance(page, int) else None,
            metadata={
                key: item[key]
                for key in ("bbox", "score")
                if key in item
            },
        )

    @staticmethod
    def _page_count(items: list[dict[str, Any]]) -> int | None:
        pages = [
            int(item["page_idx"])
            for item in items
            if isinstance(item.get("page_idx"), int)
        ]
        return max(pages) + 1 if pages else None


class FallbackDocumentParser:
    """Try the primary adapter, falling back only on configured failures."""

    def __init__(
        self,
        primary: DocumentParser,
        fallback: DocumentParser,
    ) -> None:
        self.primary = primary
        self.fallback = fallback

    async def parse(
        self,
        *,
        filename: str,
        content: bytes,
        content_type: str,
    ) -> ParsedDocument:
        try:
            return await self.primary.parse(
                filename=filename,
                content=content,
                content_type=content_type,
            )
        except DocumentParseError as exc:
            result = await self.fallback.parse(
                filename=filename,
                content=content,
                content_type=content_type,
            )
            return ParsedDocument(
                filename=result.filename,
                parser=result.parser,
                blocks=result.blocks,
                page_count=result.page_count,
                metadata={
                    **result.metadata,
                    "fallback_from": getattr(
                        self.primary, "name", "primary"
                    ),
                    "fallback_reason": str(exc),
                },
            )


@dataclass(frozen=True)
class QualityIssue:
    code: str
    severity: str
    message: str


@dataclass(frozen=True)
class DocumentQualityReport:
    score: int
    passed: bool
    issues: tuple[QualityIssue, ...]
    metrics: dict[str, Any]

    def as_dict(self) -> dict[str, Any]:
        return {
            "score": self.score,
            "passed": self.passed,
            "issues": [
                {
                    "code": issue.code,
                    "severity": issue.severity,
                    "message": issue.message,
                }
                for issue in self.issues
            ],
            "metrics": self.metrics,
        }


class DocumentQualityGate:
    """Deterministic pre-index checks with explainable scoring."""

    def __init__(
        self,
        *,
        minimum_score: int = 60,
        minimum_characters: int = 20,
        maximum_replacement_ratio: float = 0.02,
        maximum_duplicate_ratio: float = 0.5,
    ) -> None:
        self.minimum_score = minimum_score
        self.minimum_characters = minimum_characters
        self.maximum_replacement_ratio = maximum_replacement_ratio
        self.maximum_duplicate_ratio = maximum_duplicate_ratio

    def inspect(self, document: ParsedDocument) -> DocumentQualityReport:
        text = document.text.strip()
        issues: list[QualityIssue] = []
        score = 100
        replacement_ratio = text.count("\ufffd") / max(len(text), 1)
        lines = [
            re.sub(r"\s+", " ", line).strip()
            for line in text.splitlines()
            if line.strip()
        ]
        duplicate_ratio = (
            1 - (len(set(lines)) / len(lines)) if lines else 0.0
        )
        if len(text) < self.minimum_characters:
            score -= 70
            issues.append(
                QualityIssue(
                    "insufficient_text",
                    "error",
                    "可提取文本过少，可能是扫描件、空白文档或解析失败。",
                )
            )
        if replacement_ratio > self.maximum_replacement_ratio:
            score -= 50
            issues.append(
                QualityIssue(
                    "encoding_corruption",
                    "error",
                    "乱码替代字符比例过高。",
                )
            )
        if duplicate_ratio > self.maximum_duplicate_ratio:
            score -= 25
            issues.append(
                QualityIssue(
                    "duplicate_content",
                    "warning",
                    "重复文本比例过高，可能包含页眉页脚或解析重复。",
                )
            )
        empty_blocks = sum(not block.text.strip() for block in document.blocks)
        if document.blocks and empty_blocks / len(document.blocks) > 0.5:
            score -= 30
            issues.append(
                QualityIssue(
                    "empty_blocks",
                    "warning",
                    "超过一半的内容块为空。",
                )
            )
        score = max(0, score)
        passed = score >= self.minimum_score and not any(
            issue.severity == "error" for issue in issues
        )
        return DocumentQualityReport(
            score=score,
            passed=passed,
            issues=tuple(issues),
            metrics={
                "characters": len(text),
                "blocks": len(document.blocks),
                "pages": document.page_count,
                "replacement_ratio": round(replacement_ratio, 4),
                "duplicate_ratio": round(duplicate_ratio, 4),
            },
        )
